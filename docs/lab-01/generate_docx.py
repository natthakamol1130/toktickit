import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_placeholder_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=240, bottom=240, left=200, right=200)
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="dashed" w:sz="12" w:space="0" w:color="94A3B8"/>
            <w:left w:val="dashed" w:sz="12" w:space="0" w:color="94A3B8"/>
            <w:bottom w:val="dashed" w:sz="12" w:space="0" w:color="94A3B8"/>
            <w:right w:val="dashed" w:sz="12" w:space="0" w:color="94A3B8"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    run_icon = p.add_run("📷 ")
    run_icon.font.size = Pt(14)
    
    run = p.add_run(text)
    run.font.name = "Sarabun"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph() # spacing

def create_report_docx():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Sarabun'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("CPE 334 — Lab 1 Submission Report")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 118, 110)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("TokTickIT Full-Stack Hello World Starter & Git Engineering Workflow")
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    # Metadata Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    meta_data = [
        [("Student Name:", " Natthakamol Mornparn"), ("Student ID:", " 67070505215")],
        [("GitHub Account:", " @natthakamol1130"), ("Peer Reviewer:", " Chanya (6707051058, @chanya06)")]
    ]
    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(3.25)
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            lbl, val = meta_data[row_idx][col_idx]
            r1 = p.add_run(lbl)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(15, 118, 110)
            r2 = p.add_run(val)
            r2.font.color.rgb = RGBColor(30, 41, 59)
            
    doc.add_paragraph() # space

    # -------------------------------------------------------------
    # PART 1
    # -------------------------------------------------------------
    p_h1 = doc.add_heading(level=1)
    p_h1.paragraph_format.space_before = Pt(14)
    p_h1.paragraph_format.space_after = Pt(8)
    run_h1 = p_h1.add_run("Answer Part 1: Git Use with Engineering Workflow")
    run_h1.font.color.rgb = RGBColor(15, 118, 110)
    run_h1.font.size = Pt(14)

    # 1. URLs List
    p_u = doc.add_paragraph()
    p_u.add_run("1. Repository & Project URLs List").font.bold = True
    
    urls = [
        ("GitHub Repository:", "https://github.com/natthakamol1130/toktickit"),
        ("GitHub Project (Kanban):", "https://github.com/users/natthakamol1130/projects"),
        ("Issue #1 (Project foundation):", "https://github.com/natthakamol1130/toktickit/issues/1"),
        ("Issue #2 (API Health check):", "https://github.com/natthakamol1130/toktickit/issues/2"),
        ("Issue #3 (Create & seed categories):", "https://github.com/natthakamol1130/toktickit/issues/3"),
        ("Issue #4 (Display category list UI):", "https://github.com/natthakamol1130/toktickit/issues/4"),
        ("PR #5 (Feature 1 → lab1-staging):", "https://github.com/natthakamol1130/toktickit/pull/5"),
        ("PR #6 (Feature 2 → lab1-staging):", "https://github.com/natthakamol1130/toktickit/pull/6"),
        ("PR #7 (Feature 3 → lab1-staging):", "https://github.com/natthakamol1130/toktickit/pull/7"),
        ("PR #8 (Feature 4 → lab1-staging):", "https://github.com/natthakamol1130/toktickit/pull/8"),
        ("PR #9 (Release lab1-staging → main):", "https://github.com/natthakamol1130/toktickit/pull/9"),
    ]
    for label, link in urls:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{label} ")
        r1.font.bold = True
        r2 = p.add_run(link)
        r2.font.color.rgb = RGBColor(15, 118, 110)
        r2.font.underline = True

    # 2. Project Board
    doc.add_paragraph()
    p_k = doc.add_paragraph()
    p_k.add_run("2. GitHub Project Board Evidence (Kanban)").font.bold = True
    
    img_kanban = r"c:\Users\Windows\OneDrive\kmutt\Lab1_Starter_Scaffold\toktickit\docs\lab-01\images\kanban_board.png"
    if os.path.exists(img_kanban):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(8)
        run_img = p_img.add_run()
        run_img.add_picture(img_kanban, width=Inches(6.5))
    else:
        add_placeholder_box(doc, "[ แนบภาพ Screenshot: GitHub Project Kanban Board ที่ทุก Issue #1, #2, #3, #4 อยู่ในสถานะ 'Done' ]")

    # 3. Git Log Graph
    p_g = doc.add_paragraph()
    p_g.add_run("3. Git Workflow Evidence (git log graph on main)").font.bold = True
    
    p_code = doc.add_paragraph()
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(6)
    code_text = (
        "*   c03027e Merge pull request #9 from natthakamol1130/lab1-staging\n"
        "|\\  \n"
        "| *   ae40b84 Merge pull request #8 from natthakamol1130/feature/4-category-list\n"
        "| |\\  \n"
        "| | * 50e7cec docs: update peer review record in reviewer.md (Issue #4)\n"
        "| | * 6f90e97 feat: handle empty categories edge case gracefully in UI and tests\n"
        "| | * b7e2d9b docs: document AI prompts and reflection for Lab 1 (Issue #4)\n"
        "| | * 0f262cb feat: display IT request category list and system status UI (Issue #4)\n"
        "| |/  \n"
        "| *   4bc5182 Merge pull request #7 from natthakamol1130/feature/3-category-seed\n"
        "| |\\  \n"
        "| | * 79f9d51 feat: create Category model and seed 4 initial categories (Issue #3)\n"
        "| |/  \n"
        "| *   69dbaf0 Merge pull request #6 from natthakamol1130/feature/2-health-check\n"
        "| |\\  \n"
        "| | * eb256e6 feat: implement GET /api/health endpoint (Issue #2)\n"
        "| |/  \n"
        "| * f3bbc12 Merge pull request #5 from natthakamol1130/feature/1-project-foundation\n"
        "|/| \n"
        "| * 2dc4643 chore: update .gitignore with IDE and env rules from peer review\n"
        "| * 440e812 feat: set up TokTickIT project foundation (Issue #1)\n"
        "|/  \n"
        "* 101f877 Initial commit"
    )
    r_code = p_code.add_run(code_text)
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(8.5)
    img_git_graph = r"c:\Users\Windows\OneDrive\kmutt\Lab1_Starter_Scaffold\toktickit\docs\lab-01\images\git_log_graph.png"
    if os.path.exists(img_git_graph):
        p_img_g = doc.add_paragraph()
        p_img_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_g.paragraph_format.space_before = Pt(6)
        p_img_g.paragraph_format.space_after = Pt(8)
        run_img_g = p_img_g.add_run()
        run_img_g.add_picture(img_git_graph, width=Inches(6.5))
    else:
        add_placeholder_box(doc, "[ แนบภาพ Screenshot: Terminal แสดงคำสั่ง git log --oneline --graph -n 25 บน branch main ]")

    # 4. Directory Structure
    p_d = doc.add_paragraph()
    p_d.add_run("4. Repository Directory Structure Evidence").font.bold = True
    
    img_dir = r"c:\Users\Windows\OneDrive\kmutt\Lab1_Starter_Scaffold\toktickit\docs\lab-01\images\directory_structure.png"
    if os.path.exists(img_dir):
        p_img_d = doc.add_paragraph()
        p_img_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_d.paragraph_format.space_before = Pt(6)
        p_img_d.paragraph_format.space_after = Pt(8)
        run_img_d = p_img_d.add_run()
        run_img_d.add_picture(img_dir, width=Inches(3.8))
    else:
        add_placeholder_box(doc, "[ แนบภาพ Screenshot: แถบ Explorer ของ VS Code/IDE แสดงโครงสร้าง Directory ที่ครบถ้วน ]")

    # 5. .gitignore
    p_gi = doc.add_paragraph()
    p_gi.add_run("5. Rendered .gitignore Content").font.bold = True
    gi_text = (
        "# dependencies\nnode_modules/\n\n"
        "# env & secrets\n.env\n*.env\n.env.local\n.env.*.local\n!.env.example\n\n"
        "# build output\ndist/\nbuild/\n\n"
        "# prisma\nserver/prisma/*.db\n\n"
        "# IDE & Editor\n.vscode/\n.idea/\n*.swp\n*.swo\n\n"
        "# logs & OS\n*.log\n.DS_Store\nThumbs.db"
    )
    p_gi_c = doc.add_paragraph()
    r_gi = p_gi_c.add_run(gi_text)
    r_gi.font.name = "Consolas"
    r_gi.font.size = Pt(8.5)

    # 6. Peer review record
    p_pr = doc.add_paragraph()
    p_pr.add_run("6. Rendered docs/lab-01/reviewer.md (Peer Review Record)").font.bold = True
    
    p_pr_auth = doc.add_paragraph()
    p_pr_auth.add_run("Author: Natthakamol Mornparn — 67070505215 — GitHub: @natthakamol1130\n").font.bold = True
    p_pr_auth.add_run("Peer reviewer: Chanya — 6707051058 — GitHub: @chanya06").font.bold = True
    
    # Table of PRs
    t_pr = doc.add_table(rows=5, cols=3)
    t_pr.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_pr.autofit = False
    headers = ["PR #", "Feature Branch", "Reviewer Verdict"]
    for i, h in enumerate(headers):
        cell = t_pr.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
    
    pr_rows = [
        ("#5", "feature/1-project-foundation", "Approved"),
        ("#6", "feature/2-health-check", "Approved"),
        ("#7", "feature/3-category-seed", "Approved"),
        ("#8", "feature/4-category-list", "Approved")
    ]
    for row_idx, data in enumerate(pr_rows):
        row = t_pr.rows[row_idx+1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            if col_idx == 2:
                r.font.bold = True
                r.font.color.rgb = RGBColor(21, 128, 61)
                
    p_rev_c = doc.add_paragraph()
    p_rev_c.paragraph_format.space_before = Pt(6)
    p_rev_c.add_run("Reviewer comment I received:\n").font.bold = True
    r_rc = p_rev_c.add_run("“โดยรวมทำได้ค่อนข้างครบตาม requirement ทั้ง API, การดึง category มาแสดงใน UI และมีการรองรับ Loading กับ Error state ด้วย แนะนำเพิ่มเติมว่าอาจลองทดสอบกรณี API ตอบข้อมูลว่าง หรือไม่มี category เพื่อดูว่า UI จะแสดงผลอย่างไร จะช่วยให้รองรับ edge case ได้ครบขึ้น”\n\n")
    r_rc.font.italic = True
    p_rev_c.add_run("How I responded:\n").font.bold = True
    r_rr = p_rev_c.add_run("“ขอบคุณสำหรับคำแนะนำ ได้อัปเดตไฟล์ client/src/App.tsx ให้รองรับกรณี categories เป็นลิสต์ว่าง โดยแสดงผลข้อความ 'No categories found.' พร้อมเพิ่ม Unit Test ตรวจสอบ Edge case นี้ใน App.test.tsx และ push ขึ้น GitHub เรียบร้อยแล้ว”\n\n")
    r_rr.font.italic = True
    p_rev_c.add_run("Pull Requests I reviewed for my partner:\n").font.bold = True
    r_rm = p_rev_c.add_run("“ตรวจสอบโค้ดและเอกสารใน PR (Issue 4) ทั้งหมดแล้วครับ พัฒนาได้สมบูรณ์แบบมาก Endpoint GET /api/categories ดึงข้อมูลและเรียงตาม id ถูกต้อง, หน้าจอ React UI และ Unit Tests ครอบคลุมทั้ง Online, Offline และ Loading, บันทึกผลเทสต์ใน tests.md และ Prompts ใน ai_use.md ได้ละเอียดครบถ้วน Approved!”\n\n")
    r_rm.font.italic = True
    p_rev_c.add_run("Partner's response:\n").font.bold = True
    r_pr = p_rev_c.add_run("“ขอบคุณสำหรับ Review และคำแนะนำครับ ได้ตรวจสอบผลการรันเทสต์และเตรียม Merge เข้า lab1-staging เรียบร้อยครับ”")
    r_pr.font.italic = True

    # -------------------------------------------------------------
    # PART 2
    # -------------------------------------------------------------
    p_h2 = doc.add_heading(level=1)
    p_h2.paragraph_format.space_before = Pt(18)
    p_h2.paragraph_format.space_after = Pt(8)
    run_h2 = p_h2.add_run("Answer Part 2: Tests")
    run_h2.font.color.rgb = RGBColor(15, 118, 110)
    run_h2.font.size = Pt(14)

    p_t_lbl = doc.add_paragraph()
    p_t_lbl.add_run("1. Test Plan Table (docs/lab-01/tests.md)").font.bold = True
    
    t_tests = doc.add_table(rows=7, cols=5)
    t_tests.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tests.autofit = False
    test_headers = ["Test ID", "File Path", "Tool", "Test Description", "Result"]
    for i, h in enumerate(test_headers):
        cell = t_tests.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        
    test_rows = [
        ("API-01", "server/tests/lab-01/health.test.ts", "Supertest", "Health endpoint returns 200 and expected JSON status=ok", "Pass"),
        ("API-02", "server/tests/lab-01/categories.test.ts", "Supertest", "Categories endpoint returns 4 seeded categories in ID order", "Pass"),
        ("UI-01", "client/tests/lab-01/App.test.tsx", "Vitest", "TokTickIT heading renders properly", "Pass"),
        ("UI-02", "client/tests/lab-01/App.test.tsx", "Vitest", "Success state shows Online + 4 categories list", "Pass"),
        ("UI-03", "client/tests/lab-01/App.test.tsx", "Vitest", "Empty categories shows 'No categories found.' fallback", "Pass"),
        ("UI-04", "client/tests/lab-01/App.test.tsx", "Vitest", "API failure displays Offline + useful error message", "Pass"),
    ]
    for row_idx, data in enumerate(test_rows):
        row = t_tests.rows[row_idx+1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            if col_idx == 4:
                r.font.bold = True
                r.font.color.rgb = RGBColor(21, 128, 61)

    p_t_ev = doc.add_paragraph()
    p_t_ev.paragraph_format.space_before = Pt(10)
    p_t_ev.add_run("2. Terminal Output Evidence on main branch").font.bold = True
    
    p_t_out = doc.add_paragraph()
    p_t_out.paragraph_format.space_before = Pt(2)
    p_t_out.paragraph_format.space_after = Pt(6)
    out_text = (
        "> toktickit-server@1.0.0 test\n"
        "> vitest run\n\n"
        " ✓ tests/lab-01/health.test.ts (1 test) 20ms\n"
        " ✓ tests/lab-01/categories.test.ts (1 test) 194ms\n\n"
        " Test Files  2 passed (2)\n"
        "      Tests  2 passed (2)\n"
        "   Duration  10.68s\n\n"
        "--------------------------------------------------\n\n"
        "> toktickit-client@1.0.0 test\n"
        "> vitest run\n\n"
        " ✓ tests/lab-01/App.test.tsx (4 tests) 249ms\n\n"
        " Test Files  1 passed (1)\n"
        "      Tests  4 passed (4)\n"
        "   Duration  35.21s"
    )
    r_out = p_t_out.add_run(out_text)
    r_out.font.name = "Consolas"
    r_out.font.size = Pt(8.5)
    
    add_placeholder_box(doc, "[ แนบภาพ Screenshot: Terminal แสดงผลการรัน npm test ทั้งฝั่ง Server และ Client ผ่าน 100% บน branch main ]")

    # -------------------------------------------------------------
    # PART 3
    # -------------------------------------------------------------
    p_h3 = doc.add_heading(level=1)
    p_h3.paragraph_format.space_before = Pt(18)
    p_h3.paragraph_format.space_after = Pt(8)
    run_h3 = p_h3.add_run("Answer Part 3: AI Use and Reflection")
    run_h3.font.color.rgb = RGBColor(15, 118, 110)
    run_h3.font.size = Pt(14)

    p_ai_m = doc.add_paragraph()
    p_ai_m.add_run("LLM/agent used: ").font.bold = True
    p_ai_m.add_run("Antigravity (Gemini 3.6 Flash / Medium Thinking)")

    t_ai = doc.add_table(rows=9, cols=4)
    t_ai.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_ai.autofit = False
    ai_headers = ["#", "Prompt Name", "Actual Prompt Text (Summarised)", "What I did with the result"]
    for i, h in enumerate(ai_headers):
        cell = t_ai.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        
    ai_rows = [
        ("1", "Plan Lab 1 Implementation", "Read the enclosed TokTickIT Lab 1 requirements. Summarize the four GitHub Issues, their dependencies, required outputs, and required automated tests. Propose an implementation order, but do not write code yet.", "Reviewed the proposed implementation steps and approved the execution plan."),
        ("2", "Set Up Full-Stack Project", "Setup the TokTickIT project tech stack as given in Lab 1 using React, TypeScript, Vite, and Bootstrap for the frontend, and Node.js, Express, and TypeScript for the backend. Configure PostgreSQL and Prisma. Use the required folder structure.", "Verified dependencies, initialized git repository and branch structure, and created PR #5 to lab1-staging."),
        ("3", "Implement Health Check", "Add GET /api/health to the existing Express backend returning HTTP 200 with { status: \"ok\", service: \"TokTickIT API\" } and verify with Supertest.", "Verified test passed 100%, committed changes to feature/2-health-check, and created PR #6."),
        ("4", "Implement Category DB Feature", "Define Category model in Prisma schema, run migration, and implement idempotent database seeding using upsert for the 4 required categories.", "Ran Prisma migrations against PostgreSQL, verified seed idempotency, and opened PR #7."),
        ("5", "Implement Category List API", "Implement GET /api/categories endpoint in Express backend returning categories in predictable ID order from PostgreSQL via Prisma.", "Verified endpoint with Supertest in categories.test.ts."),
        ("6", "Build and Test Check System UI", "Create a Bootstrap-based page with [Check System] button. When clicked, show a loading state, call backend APIs, and render status badge and category list.", "Implemented App.tsx and api.ts, connected frontend with backend API."),
        ("7", "Write Client Unit Tests", "Write Vitest test suite for App component covering heading, success state, offline error state, and empty categories edge case.", "Ran Vitest in client, confirming all 4 test cases passed."),
        ("8", "Review Peer PRs and Finalize Lab 1", "Review peer partners' pull requests for Issues 1, 2, 3, and 4 with constructive feedback and document peer review records.", "Provided review comments, suggested improvements (.gitignore, seed idempotency), and documented review records in reviewer.md."),
    ]
    for row_idx, data in enumerate(ai_rows):
        row = t_ai.rows[row_idx+1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(text)

    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_before = Pt(8)
    p_ref.add_run("Reflection on Experience Improving Prompts:\n").font.bold = True
    r_rf = p_ref.add_run("“Providing structured prompt contexts with explicit acceptance criteria, expected JSON response formats, and clear Git branch instructions allowed the agent to implement endpoints and test assertions accurately in one shot. When handling Windows PowerShell syntax nuances and database connection configurations, detailing the local environment parameters ensured that database migrations, seed scripts, and test suites executed reliably without errors.”")
    r_rf.font.italic = True

    # -------------------------------------------------------------
    # PART 4
    # -------------------------------------------------------------
    p_h4 = doc.add_heading(level=1)
    p_h4.paragraph_format.space_before = Pt(18)
    p_h4.paragraph_format.space_after = Pt(8)
    run_h4 = p_h4.add_run("Answer Part 4: App Demo")
    run_h4.font.color.rgb = RGBColor(15, 118, 110)
    run_h4.font.size = Pt(14)

    p_a1 = doc.add_paragraph()
    p_a1.add_run("1. Initial Screen State").font.bold = True
    p_a1_desc = doc.add_paragraph()
    p_a1_desc.add_run("หน้าจอเริ่มต้นเมื่อเปิด Web Application มีข้อความหัวข้อ TokTickIT IT Service Desk และปุ่ม [Check System]")
    add_placeholder_box(doc, "[ แนบภาพ Screenshot: หน้าจอเริ่มต้นแสดงหัวข้อ TokTickIT และปุ่ม Check System ]")

    p_a2 = doc.add_paragraph()
    p_a2.add_run("2. Success Case (Backend Online & Database Connected)").font.bold = True
    p_a2_desc = doc.add_paragraph()
    p_a2_desc.add_run("เมื่อกดปุ่ม [Check System] โดยเปิด Backend Server และ Database ปกติ ระบบแสดงผลสถานะ System Status: Online พร้อมรายการ 4 หมวดหมู่ (1. Account and Access, 2. Hardware, 3. Software, 4. Network)")
    add_placeholder_box(doc, "[ แนบภาพ Screenshot: หน้าจอแสดง System Status: Online พร้อมรายการ 4 หมวดหมู่ ]")

    p_a3 = doc.add_paragraph()
    p_a3.add_run("3. Failure Case (Backend Offline / API Unavailable)").font.bold = True
    p_a3_desc = doc.add_paragraph()
    p_a3_desc.add_run("เมื่อกดปุ่ม [Check System] ในขณะที่ Backend Server หรือ Database ปิดอยู่ ระบบแสดงสถานะ System Status: Offline พร้อมข้อความแจ้งเตือน Unable to connect to TokTickIT API")
    add_placeholder_box(doc, "[ แนบภาพ Screenshot: หน้าจอแสดง System Status: Offline พร้อมข้อความ Unable to connect to TokTickIT API ]")

    output_path = r"c:\Users\Windows\OneDrive\kmutt\Lab1_Starter_Scaffold\toktickit\docs\lab-01\LAB1_SUBMISSION_REPORT.docx"
    output_ready = r"c:\Users\Windows\OneDrive\kmutt\Lab1_Starter_Scaffold\toktickit\docs\lab-01\LAB1_SUBMISSION_REPORT_READY.docx"
    doc.save(output_ready)
    print(f"Successfully generated docx at: {output_ready}")
    try:
        doc.save(output_path)
        print(f"Successfully generated docx at: {output_path}")
    except Exception as e:
        print(f"Note: {output_path} is currently open in Word. Saved to {output_ready} instead.")

if __name__ == "__main__":
    create_report_docx()
